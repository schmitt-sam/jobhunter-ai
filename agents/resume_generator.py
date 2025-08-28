import os
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
import re

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def sanitize_filename(name: str) -> str:
    return re.sub(r"[^\w\-_. ]", "_", name.strip())

def generate_tailored_resume(job_description: str, bullet_points: list[str]) -> str:
    prompt = f"""
You are a resume optimization assistant. Given the job description below and a large set of resume bullet points, generate a tailored resume that:

1. Prioritizes the most relevant experience for this job.
2. Includes strong, quantifiable results where available.
3. Organizes the resume with clear section headings.
4. Uses professional, concise language.

--- Job Description ---
{job_description}

--- Available Resume Bullet Points ---
{chr(10).join(bullet_points)}

--- Output Format ---
Return a full resume in professional markdown or plain text with headings (e.g., Summary, Skills, Experience, Education).
"""

    # Minimal input sanity logs
    print(f"Job description length: {len(job_description.strip())}")
    print(f"Bullet points count: {len(bullet_points)}")

    print("Generating resume with gpt-4o-mini...")
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=1,
        max_tokens=1500,
    )
    content = (response.choices[0].message.content or "").strip()
    if not content:
        raise RuntimeError(
            "Empty resume from model. Check API key, model name, and inputs."
        )
    return content

def save_as_docx(resume_text: str, output_path: str):
    if not resume_text or not resume_text.strip():
        raise ValueError("Refusing to save empty resume text to DOCX.")

    doc = Document()
    for line in resume_text.splitlines():
        if line.strip().startswith("##") or line.strip().endswith(":"):
            doc.add_heading(line.strip("# ").rstrip(":"), level=2)
        elif line.strip():
            doc.add_paragraph(line.strip())
    doc.save(output_path)
